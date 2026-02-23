# ai_reporter/reporter.py
# Ollama Cloud AI Commuter Report – MBTA Red Line.
# Fetches API data, formats for Ollama (token-optimized), queries Ollama Cloud,
# and writes .docx report. Prompt is parameterized by dep/arr station names.

# 0. Setup #################################################################

import json
import os
import requests
import pandas as pd
from pathlib import Path
from datetime import datetime, timezone
from zoneinfo import ZoneInfo

EASTERN = ZoneInfo("America/New_York")

from dotenv import load_dotenv
from docx import Document

## 0.1 Load environment #####################################################

_app_dir = Path(__file__).resolve().parent.parent
_repo_root = _app_dir.parent
# Prefer repo-level .env so one shared file works from any run context
_env_paths = [_repo_root / ".env", _app_dir / ".env"]
for _p in _env_paths:
    if _p.exists():
        load_dotenv(_p)
        break
else:
    load_dotenv()

OLLAMA_API_KEY = os.getenv("OLLAMA_API_KEY")
OLLAMA_CHAT_URL = "https://ollama.com/api/chat"
OLLAMA_MODEL = "gpt-oss:20b-cloud"

# 1. Prompt template (parameterized) ########################################


def get_report_prompt(departure_station_name: str, arrival_station_name: str) -> str:
    """
    Build the Ollama prompt with explicit section order and table format.
    Use station names from the data; inject dep/arr for report context.
    """
    return f"""Create a morning commute report for the Boston MBTA Red Line based on the JSON data below.

**Section order (follow exactly):**
1. Quick Takeaway – one or two sentences at the top with the essential summary (e.g., delays, on-time status, key alert).
2. Service Alerts – bullet points only; use station names from the data; do not invent stops.
3. Train Rollup Table – outbound trains from {departure_station_name}.
4. Brief summary – estimated time to {arrival_station_name} if available.

**Table format** (markdown table with these columns in this order):
| Train ID | Destination | Scheduled Departure | Estimated Departure | On Time Status | Est. Time to {arrival_station_name} |

**Constraints:**
- Maximum half page.
- Put the Quick Takeaway at the very top.
- Friendly, conversational tone.
- Use station names from the data only; do not invent stops."""


# 2. Data preparation (unfiltered, for Ollama) #############################


def _build_included_lookup(included_list: list) -> dict:
    lookup = {}
    for inc in included_list or []:
        key = (inc.get("type"), inc.get("id"))
        lookup[key] = inc
    return lookup


def build_alerts_df(alerts_response: dict) -> pd.DataFrame:
    """Build alerts DataFrame. Columns: id, severity, header, short_header, description, effect."""
    if alerts_response.get("error"):
        return pd.DataFrame()
    data = alerts_response.get("data", [])
    rows = []
    for item in data:
        attrs = item.get("attributes", {})
        rows.append({
            "id": item.get("id"),
            "severity": attrs.get("severity"),
            "header": attrs.get("header") or "",
            "short_header": attrs.get("short_header") or "",
            "description": attrs.get("description") or "",
            "effect": attrs.get("effect"),
        })
    if not rows:
        return pd.DataFrame()
    return pd.DataFrame(rows)


def build_predictions_df(predictions_response: dict) -> pd.DataFrame:
    """Flatten predictions into one row per prediction with schedule, trip, stop, vehicle."""
    if predictions_response.get("error"):
        return pd.DataFrame()
    payload = predictions_response
    included = _build_included_lookup(payload.get("included", []))
    rows = []
    for item in payload.get("data", []):
        attrs = item.get("attributes", {})
        rels = item.get("relationships", {})
        schedule_id = (rels.get("schedule") or {}).get("data")
        schedule_id = schedule_id.get("id") if isinstance(schedule_id, dict) else None
        schedule = included.get(("schedule", schedule_id)) if schedule_id else {}
        sched_attrs = schedule.get("attributes", {})
        trip_id = (rels.get("trip") or {}).get("data")
        trip_id = trip_id.get("id") if isinstance(trip_id, dict) else None
        trip = included.get(("trip", trip_id)) if trip_id else {}
        trip_attrs = trip.get("attributes", {})
        stop_id = (rels.get("stop") or {}).get("data")
        stop_id = stop_id.get("id") if isinstance(stop_id, dict) else None
        stop = included.get(("stop", stop_id)) if stop_id else {}
        stop_attrs = stop.get("attributes", {})
        vehicle_ref = (rels.get("vehicle") or {}).get("data")
        vehicle_id = vehicle_ref.get("id") if isinstance(vehicle_ref, dict) else None
        rows.append({
            "direction_id": attrs.get("direction_id"),
            "departure_time": attrs.get("departure_time"),
            "arrival_time": attrs.get("arrival_time"),
            "schedule_departure_time": sched_attrs.get("departure_time") or sched_attrs.get("departure"),
            "schedule_arrival_time": sched_attrs.get("arrival_time") or sched_attrs.get("arrival"),
            "trip_headsign": trip_attrs.get("headsign", ""),
            "stop_name": stop_attrs.get("name", ""),
            "vehicle_id": vehicle_id,
        })
    if not rows:
        return pd.DataFrame()
    return pd.DataFrame(rows)


def build_vehicles_df(vehicles_response: dict) -> pd.DataFrame:
    """Flatten vehicles with trip and stop from included."""
    if vehicles_response.get("error"):
        return pd.DataFrame()
    payload = vehicles_response
    included = _build_included_lookup(payload.get("included", []))
    rows = []
    for item in payload.get("data", []):
        attrs = item.get("attributes", {})
        rels = item.get("relationships", {})
        stop_ref = (rels.get("stop") or {}).get("data")
        stop_id = stop_ref.get("id") if isinstance(stop_ref, dict) else None
        stop_res = included.get(("stop", stop_id)) if stop_id else {}
        stop_name = (stop_res.get("attributes") or {}).get("name", "") if stop_res else ""
        trip_ref = (rels.get("trip") or {}).get("data")
        trip_id = trip_ref.get("id") if isinstance(trip_ref, dict) else None
        trip = included.get(("trip", trip_id)) if trip_id else {}
        trip_attrs = trip.get("attributes", {}) if trip else {}
        lat = attrs.get("latitude")
        if lat is None and isinstance(attrs.get("position"), dict):
            lat = attrs["position"].get("latitude")
        lon = attrs.get("longitude")
        if lon is None and isinstance(attrs.get("position"), dict):
            lon = attrs["position"].get("longitude")
        rows.append({
            "vehicle_id": item.get("id"),
            "latitude": lat,
            "longitude": lon,
            "current_stop_name": stop_name,
            "trip_headsign": trip_attrs.get("headsign", ""),
        })
    if not rows:
        return pd.DataFrame()
    return pd.DataFrame(rows)


# 3. Token-optimized format for Ollama ######################################


def format_data_for_ollama_compact(
    df_alerts: pd.DataFrame,
    df_predictions_dep: pd.DataFrame,
    df_predictions_arr: pd.DataFrame,
    df_vehicles: pd.DataFrame,
    dep_label: str = "Departure stop",
    arr_label: str = "Arrival stop",
) -> str:
    """
    Compact format: truncate alerts, drop low-value cols, limit rows.
    Station-scoped predictions; time-window already applied in build_predictions_df.
    """
    alerts_cols = ["id", "header", "short_header", "description", "effect"]
    alerts_sub = df_alerts[[c for c in alerts_cols if c in df_alerts.columns]] if not df_alerts.empty else pd.DataFrame()
    pred_cols = ["direction_id", "departure_time", "arrival_time", "schedule_departure_time", "schedule_arrival_time", "trip_headsign", "stop_name", "vehicle_id"]
    pred_dep = df_predictions_dep[[c for c in pred_cols if c in df_predictions_dep.columns]] if not df_predictions_dep.empty else pd.DataFrame()
    pred_arr = df_predictions_arr[[c for c in pred_cols if c in df_predictions_arr.columns]] if not df_predictions_arr.empty else pd.DataFrame()
    veh_cols = ["vehicle_id", "current_stop_name", "trip_headsign"]
    veh_sub = df_vehicles[[c for c in veh_cols if c in df_vehicles.columns]] if not df_vehicles.empty else pd.DataFrame()
    sections = []
    sections.append("## Alerts\n" + (alerts_sub.to_json(orient="records", date_format="iso") if not alerts_sub.empty else "[]"))
    sections.append(f"## Predictions at {dep_label}\n" + (pred_dep.to_json(orient="records", date_format="iso") if not pred_dep.empty else "[]"))
    sections.append(f"## Predictions at {arr_label}\n" + (pred_arr.to_json(orient="records", date_format="iso") if not pred_arr.empty else "[]"))
    sections.append("## Vehicles\n" + (veh_sub.to_json(orient="records", date_format="iso") if not veh_sub.empty else "[]"))
    return "\n\n".join(sections)


# 4. Ollama Cloud query #####################################################


def query_ollama_cloud(user_content: str) -> str:
    """Send user_content to Ollama Cloud chat API and return the assistant message."""
    if not OLLAMA_API_KEY:
        raise ValueError("OLLAMA_API_KEY is not set. Add it to .env.")
    headers = {"Authorization": f"Bearer {OLLAMA_API_KEY}", "Content-Type": "application/json"}
    body = {"model": OLLAMA_MODEL, "messages": [{"role": "user", "content": user_content}], "stream": False}
    try:
        response = requests.post(OLLAMA_CHAT_URL, headers=headers, json=body, timeout=60)
    except requests.RequestException as e:
        raise ValueError(f"Ollama request failed: {e}") from e
    if response.status_code != 200:
        raise ValueError(f"Ollama API error: status {response.status_code}, body={response.text[:500]}")
    try:
        result = response.json()
    except Exception as e:
        raise ValueError(f"Ollama response not valid JSON: {e}") from e
    message = result.get("message")
    if not message or "content" not in message:
        raise ValueError("Ollama response missing message.content")
    return message["content"]


# 5. Write report to .docx (YYYY.MM.DD_HHMM naming) ########################


def _add_paragraph_with_inline_format(doc: Document, line: str, style: str | None = None) -> None:
    """Add a paragraph, converting **bold** and *italic* to native Word formatting."""
    p = doc.add_paragraph(style=style)
    if not line:
        return
    remaining = line
    while remaining:
        star2 = remaining.find("**")
        star1 = remaining.find("*")
        if star2 >= 0 and (star1 < 0 or star2 <= star1):
            before = remaining[:star2]
            after = remaining[star2 + 2:]
            end_bold = after.find("**")
            if end_bold >= 0:
                p.add_run(before)
                r = p.add_run(after[:end_bold])
                r.bold = True
                remaining = after[end_bold + 2:]
            else:
                p.add_run(remaining)
                break
        elif star1 >= 0:
            before = remaining[:star1]
            after = remaining[star1 + 1:]
            end_italic = after.find("*")
            if end_italic >= 0:
                p.add_run(before)
                r = p.add_run(after[:end_italic])
                r.italic = True
                remaining = after[end_italic + 1:]
            else:
                p.add_run(remaining)
                break
        else:
            p.add_run(remaining)
            break


def _is_markdown_table_row(line: str) -> bool:
    s = line.strip()
    return len(s) >= 2 and s.startswith("|") and s.endswith("|")


def _is_markdown_table_separator(line: str) -> bool:
    s = line.strip()
    if not _is_markdown_table_row(s):
        return False
    inner = s[1:-1].strip()
    return all(c in "|-: " for c in inner)


def _parse_markdown_table_lines(table_lines: list[str]) -> list[list[str]]:
    rows = []
    for line in table_lines:
        stripped = line.strip()
        if not stripped or _is_markdown_table_separator(stripped):
            continue
        cells = [c.strip() for c in stripped.split("|") if c is not None]
        if cells and cells[0] == "" and cells[-1] == "":
            cells = cells[1:-1]
        elif cells and cells[0] == "":
            cells = cells[1:]
        elif cells and cells[-1] == "":
            cells = cells[:-1]
        if cells:
            rows.append(cells)
    return rows


def _add_markdown_table_to_docx(doc: Document, table_lines: list[str]) -> None:
    rows = _parse_markdown_table_lines(table_lines)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    nrows = len(rows)
    table = doc.add_table(rows=nrows, cols=ncols)
    table.style = "Table Grid"
    for i, row_cells in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_cells):
            if j < ncols:
                cell = row.cells[j]
                cell.text = cell_text
                if i == 0:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.bold = True
        for j in range(len(row_cells), ncols):
            row.cells[j].text = ""
    doc.add_paragraph()


def _markdown_line_to_docx(doc: Document, line: str) -> None:
    stripped = line.strip()
    if not stripped:
        doc.add_paragraph()
        return
    if stripped.startswith("### "):
        doc.add_heading(stripped[4:].strip(), level=3)
    elif stripped.startswith("## "):
        doc.add_heading(stripped[3:].strip(), level=2)
    elif stripped.startswith("# "):
        doc.add_heading(stripped[2:].strip(), level=1)
    elif stripped.startswith("- ") or (stripped.startswith("* ") and not stripped.startswith("**")):
        bullet_text = stripped[2:].strip()
        _add_paragraph_with_inline_format(doc, bullet_text, style="List Bullet")
    else:
        _add_paragraph_with_inline_format(doc, stripped)


def _report_blocks(report_text: str) -> list[tuple[str, str | list[str]]]:
    lines = report_text.strip().split("\n")
    blocks = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if _is_markdown_table_row(line.strip()):
            table_lines = []
            while i < len(lines) and _is_markdown_table_row(lines[i].strip()):
                table_lines.append(lines[i])
                i += 1
            blocks.append(("table", table_lines))
            continue
        blocks.append(("paragraph", line))
        i += 1
    return blocks


def write_report_docx(
    report_text: str,
    output_dir: Path | None = None,
) -> Path:
    """
    Write report_text to a Word document. Naming: MBTA Red Line Commuter Report YYYY.MM.DD_HHMM.docx
    Returns the path where the file was saved.
    """
    if output_dir is None:
        output_dir = _app_dir / "reports"
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    # Use UTC then convert to Eastern (avoids Windows locale/TZ env issues)
    ts = datetime.now(timezone.utc).astimezone(EASTERN)
    fname = f"MBTA Red Line Commuter Report {ts.strftime('%Y.%m.%d')}_{ts.strftime('%H%M')}.docx"
    output_path = output_dir / fname
    doc = Document()
    doc.add_heading("MBTA Red Line – Morning Commute Overview", level=0)
    doc.add_paragraph(ts.strftime("%Y-%m-%d %H:%M %Z"))
    doc.add_paragraph()
    blocks = _report_blocks(report_text)
    for kind, content in blocks:
        if kind == "table":
            _add_markdown_table_to_docx(doc, content)
        else:
            _markdown_line_to_docx(doc, content)
    doc.save(str(output_path))
    return output_path
