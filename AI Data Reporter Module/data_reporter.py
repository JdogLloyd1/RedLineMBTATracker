# data_reporter.py
# Ollama Cloud Agent MBTA Red Line Tracker Data Reporter
# Query the MBTA API and synthesize data into a commute report (Alewife → Central Square).
# Jonathan Lloyd
#
# Fetches unfiltered API data, structures it into dataframes, sends to Ollama Cloud,
# and writes the model response to a .docx file. User fills USER_PROMPT separately.

import json
import os
import requests
import pandas as pd
from pathlib import Path
from datetime import datetime, timezone

from dotenv import load_dotenv
from docx import Document

# User fills prompt below; data will be appended automatically when the script runs.
USER_PROMPT = """
Hi Ollama! I'm a commuter on the Boston MBTA Red Line. 
I need you to help me create a morning commute report based on real-time data from the MBTA API. 
The data you are provided is in JSON format. Synthesize the data into a readable report.
The report should be written in a friendly, conversational tone. Maximum length should be half a page.
Include the following information in the report: 
- Service alerts on the Red Line
- A rollup table of data on outbound trains from Alewife:
  - Train ID
  - Destination
  - Scheduled Departure Time from Alewife
  - Estimated Departure Time from Alewife
  - On Time Status 
  - Estimated Time to Destination Central Square
"""

# 0. Setup #################################################################

## 0.1 Load environment #####################################################
# Load .env from this module's folder or project root.

_app_dir = Path(__file__).resolve().parent
_env_paths = [_app_dir / ".env", _app_dir.parent / ".env"]
for _p in _env_paths:
    if _p.exists():
        load_dotenv(_p)
        break
else:
    load_dotenv()

BASE_URL = "https://api-v3.mbta.com"
MBTA_API_KEY = os.getenv("MBTA_API_KEY")
OLLAMA_API_KEY = os.getenv("OLLAMA_API_KEY")
HEADERS = {"x-api-key": MBTA_API_KEY} if MBTA_API_KEY else {}

# Stop IDs for Alewife (origin) and Central Square (destination).
STOP_ALEWIFE = "place-alfcl"
STOP_CENTRAL_SQ = "place-cntsq"

# 1. API layer #################################################################


def _request(path: str, params: dict | None = None) -> dict:
    """
    GET request to MBTA API. Returns JSON body on success.
    On failure returns {"error": True, "message": "...", "status_code": int}.
    """
    if not MBTA_API_KEY:
        return {
            "error": True,
            "message": "MBTA_API_KEY is not set. Add it to .env.",
            "status_code": None,
        }
    url = f"{BASE_URL}{path}"
    try:
        resp = requests.get(url, headers=HEADERS, params=params or {}, timeout=15)
    except requests.RequestException as e:
        return {"error": True, "message": str(e), "status_code": None}
    if resp.status_code != 200:
        return {
            "error": True,
            "message": f"API request failed: status {resp.status_code}",
            "status_code": resp.status_code,
        }
    try:
        return resp.json()
    except Exception as e:
        return {
            "error": True,
            "message": f"Invalid JSON: {e}",
            "status_code": resp.status_code,
        }


def fetch_alerts() -> dict:
    """Fetch Red Line service alerts. Returns JSON data or error dict."""
    return _request("/alerts", params={"filter[route]": "Red"})


def fetch_predictions_at_stop(stop_id: str) -> dict:
    """Fetch predictions at a given stop for Red Line with schedule, trip, stop, vehicle."""
    return _request(
        "/predictions",
        params={
            "filter[stop]": stop_id,
            "filter[route]": "Red",
            "include": "schedule,trip,stop,vehicle",
        },
    )


def fetch_vehicles() -> dict:
    """Fetch Red Line vehicles (current stop and position) with trip and stop included."""
    return _request("/vehicles", params={"filter[route]": "Red", "include": "trip,stop"})


# 2. Data preparation (unfiltered dataframes) #####################################


def _build_included_lookup(included_list: list) -> dict:
    """Build (type, id) -> resource lookup from JSON:API included array."""
    lookup = {}
    for inc in included_list or []:
        key = (inc.get("type"), inc.get("id"))
        lookup[key] = inc
    return lookup


def build_alerts_df(alerts_response: dict) -> pd.DataFrame:
    """
    Build alerts DataFrame with full attributes; no truncation.
    Columns: id, severity, header, short_header, description, created_at, updated_at,
    active_period (JSON string), etc.
    """
    if alerts_response.get("error"):
        return pd.DataFrame()
    data = alerts_response.get("data", [])
    rows = []
    for item in data:
        attrs = item.get("attributes", {})
        active_period = attrs.get("active_period")
        rows.append({
            "id": item.get("id"),
            "severity": attrs.get("severity"),
            "header": attrs.get("header") or "",
            "short_header": attrs.get("short_header") or "",
            "description": attrs.get("description") or "",
            "created_at": attrs.get("created_at"),
            "updated_at": attrs.get("updated_at"),
            "active_period": json.dumps(active_period) if active_period is not None else "",
            "cause": attrs.get("cause"),
            "effect": attrs.get("effect"),
            "url": attrs.get("url") or "",
        })
    if not rows:
        return pd.DataFrame()
    return pd.DataFrame(rows)


def build_predictions_df(predictions_response: dict) -> pd.DataFrame:
    """
    Flatten predictions response into one row per prediction with joined schedule, trip, stop, vehicle.
    All predictions kept (both directions); no time-window filtering.
    """
    if predictions_response.get("error"):
        return pd.DataFrame()
    payload = predictions_response
    included = _build_included_lookup(payload.get("included", []))
    rows = []
    for item in payload.get("data", []):
        attrs = item.get("attributes", {})
        rels = item.get("relationships", {})
        # Resolve schedule
        schedule_id = (rels.get("schedule") or {}).get("data")
        schedule_id = schedule_id.get("id") if isinstance(schedule_id, dict) else None
        schedule = included.get(("schedule", schedule_id)) if schedule_id else {}
        sched_attrs = schedule.get("attributes", {})
        # Resolve trip
        trip_id = (rels.get("trip") or {}).get("data")
        trip_id = trip_id.get("id") if isinstance(trip_id, dict) else None
        trip = included.get(("trip", trip_id)) if trip_id else {}
        trip_attrs = trip.get("attributes", {})
        # Resolve stop
        stop_id = (rels.get("stop") or {}).get("data")
        stop_id = stop_id.get("id") if isinstance(stop_id, dict) else None
        stop = included.get(("stop", stop_id)) if stop_id else {}
        stop_attrs = stop.get("attributes", {})
        # Resolve vehicle
        vehicle_ref = (rels.get("vehicle") or {}).get("data")
        vehicle_id = vehicle_ref.get("id") if isinstance(vehicle_ref, dict) else None
        rows.append({
            "prediction_id": item.get("id"),
            "direction_id": attrs.get("direction_id"),
            "departure_time": attrs.get("departure_time"),
            "arrival_time": attrs.get("arrival_time"),
            "schedule_departure_time": sched_attrs.get("departure_time") or sched_attrs.get("departure"),
            "schedule_arrival_time": sched_attrs.get("arrival_time") or sched_attrs.get("arrival"),
            "trip_id": trip_id,
            "trip_headsign": trip_attrs.get("headsign", ""),
            "trip_direction_id": trip_attrs.get("direction_id"),
            "stop_id": stop_id,
            "stop_name": stop_attrs.get("name", ""),
            "vehicle_id": vehicle_id,
        })
    if not rows:
        return pd.DataFrame()
    return pd.DataFrame(rows)


def build_vehicles_df(vehicles_response: dict) -> pd.DataFrame:
    """
    Flatten vehicles response with trip and stop from included.
    Columns: vehicle id, latitude, longitude, bearing, current stop name, trip headsign, etc.
    """
    if vehicles_response.get("error"):
        return pd.DataFrame()
    payload = vehicles_response
    included = _build_included_lookup(payload.get("included", []))
    rows = []
    for item in payload.get("data", []):
        attrs = item.get("attributes", {})
        rels = item.get("relationships", {})
        # Current stop
        stop_ref = (rels.get("stop") or {}).get("data")
        stop_id = stop_ref.get("id") if isinstance(stop_ref, dict) else None
        stop_res = included.get(("stop", stop_id)) if stop_id else {}
        stop_name = (stop_res.get("attributes") or {}).get("name", "") if stop_res else ""
        # Trip
        trip_ref = (rels.get("trip") or {}).get("data")
        trip_id = trip_ref.get("id") if isinstance(trip_ref, dict) else None
        trip = included.get(("trip", trip_id)) if trip_id else {}
        trip_attrs = trip.get("attributes", {}) if trip else {}
        # Position (support both top-level and nested position)
        lat = attrs.get("latitude")
        if lat is None and isinstance(attrs.get("position"), dict):
            lat = attrs["position"].get("latitude")
        lon = attrs.get("longitude")
        if lon is None and isinstance(attrs.get("position"), dict):
            lon = attrs["position"].get("longitude")
        rows.append({
            "vehicle_id": item.get("id"),
            "latitude": lat,
            "longitude": lon,
            "bearing": attrs.get("bearing"),
            "current_stop_id": stop_id,
            "current_stop_name": stop_name,
            "trip_id": trip_id,
            "trip_headsign": trip_attrs.get("headsign", ""),
            "trip_direction_id": trip_attrs.get("direction_id"),
            "updated_at": attrs.get("updated_at"),
        })
    if not rows:
        return pd.DataFrame()
    return pd.DataFrame(rows)


# 3. Format data for Ollama #####################################################


def format_data_for_ollama(
    df_alerts: pd.DataFrame,
    df_predictions_alewife: pd.DataFrame,
    df_predictions_central_sq: pd.DataFrame,
    df_vehicles: pd.DataFrame,
) -> str:
    """
    Convert each DataFrame to a machine-readable string and combine with section headers.
    Uses JSON records with ISO dates for clarity.
    """
    sections = []
    sections.append("## Alerts\n" + (df_alerts.to_json(orient="records", date_format="iso") if not df_alerts.empty else "[]"))
    sections.append("## Predictions at Alewife\n" + (df_predictions_alewife.to_json(orient="records", date_format="iso") if not df_predictions_alewife.empty else "[]"))
    sections.append("## Predictions at Central Square\n" + (df_predictions_central_sq.to_json(orient="records", date_format="iso") if not df_predictions_central_sq.empty else "[]"))
    sections.append("## Vehicles\n" + (df_vehicles.to_json(orient="records", date_format="iso") if not df_vehicles.empty else "[]"))
    return "\n\n".join(sections)


# 4. Ollama Cloud query ########################################################


OLLAMA_CHAT_URL = "https://ollama.com/api/chat"
OLLAMA_MODEL = "gpt-oss:20b-cloud"


def query_ollama_cloud(user_content: str) -> str:
    """
    Send user_content to Ollama Cloud chat API and return the assistant message.
    Raises ValueError if OLLAMA_API_KEY is missing or on API error.
    """
    if not OLLAMA_API_KEY:
        raise ValueError("OLLAMA_API_KEY is not set. Add it to .env.")
    headers = {
        "Authorization": f"Bearer {OLLAMA_API_KEY}",
        "Content-Type": "application/json",
    }
    body = {
        "model": OLLAMA_MODEL,
        "messages": [{"role": "user", "content": user_content}],
        "stream": False,
    }
    try:
        response = requests.post(OLLAMA_CHAT_URL, headers=headers, json=body, timeout=60)
    except requests.RequestException as e:
        raise ValueError(f"Ollama request failed: {e}") from e
    if response.status_code != 200:
        raise ValueError(
            f"Ollama API error: status {response.status_code}, body={response.text[:500]}"
        )
    try:
        result = response.json()
    except Exception as e:
        raise ValueError(f"Ollama response not valid JSON: {e}") from e
    message = result.get("message")
    if not message or "content" not in message:
        raise ValueError("Ollama response missing message.content")
    return message["content"]


# 5. Write report to .docx ####################################################


def _add_paragraph_with_inline_format(doc: Document, line: str, style: str | None = None) -> None:
    """
    Add a paragraph, converting **bold** and *italic* to native Word formatting.
    Splits on ** and * to create runs with bold/italic set appropriately.
    """
    p = doc.add_paragraph(style=style)
    if not line:
        return
    # Parse **bold** and *italic*: process in order, splitting by ** first then *.
    remaining = line
    while remaining:
        # Find next ** or *
        star2 = remaining.find("**")
        star1 = remaining.find("*")
        # Prefer ** if it appears first and is not part of *
        if star2 >= 0 and (star1 < 0 or star2 <= star1):
            before = remaining[:star2]
            after = remaining[star2 + 2:]
            end_bold = after.find("**")
            if end_bold >= 0:
                p.add_run(before)
                r = p.add_run(after[:end_bold])
                r.bold = True
                remaining = after[end_bold + 2:]
            else:
                p.add_run(remaining)
                break
        elif star1 >= 0:
            before = remaining[:star1]
            after = remaining[star1 + 1:]
            end_italic = after.find("*")
            if end_italic >= 0:
                p.add_run(before)
                r = p.add_run(after[:end_italic])
                r.italic = True
                remaining = after[end_italic + 1:]
            else:
                p.add_run(remaining)
                break
        else:
            p.add_run(remaining)
            break


def _is_markdown_table_row(line: str) -> bool:
    """True if the line looks like a markdown table row (starts and ends with |)."""
    s = line.strip()
    return len(s) >= 2 and s.startswith("|") and s.endswith("|")


def _is_markdown_table_separator(line: str) -> bool:
    """True if the line is the markdown table separator (|---|---|)."""
    s = line.strip()
    if not _is_markdown_table_row(s):
        return False
    inner = s[1:-1].strip()
    return all(c in "|-: " for c in inner)


def _parse_markdown_table_lines(table_lines: list[str]) -> list[list[str]]:
    """
    Parse consecutive markdown table lines into a list of rows (each row is a list of cell strings).
    Skips the separator line (|---|---|). Strips cell content.
    """
    rows = []
    for line in table_lines:
        stripped = line.strip()
        if not stripped:
            continue
        if _is_markdown_table_separator(stripped):
            continue
        cells = [c.strip() for c in stripped.split("|") if c is not None]
        if cells and cells[0] == "" and cells[-1] == "":
            cells = cells[1:-1]
        elif cells and cells[0] == "":
            cells = cells[1:]
        elif cells and cells[-1] == "":
            cells = cells[:-1]
        if cells:
            rows.append(cells)
    return rows


def _add_markdown_table_to_docx(doc: Document, table_lines: list[str]) -> None:
    """
    Parse markdown table lines and append a native Word table to the document.
    First row is used as header (bold). Ensures all rows have the same column count.
    """
    rows = _parse_markdown_table_lines(table_lines)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    nrows = len(rows)
    table = doc.add_table(rows=nrows, cols=ncols)
    table.style = "Table Grid"
    for i, row_cells in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_cells):
            if j < ncols:
                cell = row.cells[j]
                cell.text = cell_text
                if i == 0:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.bold = True
        for j in range(len(row_cells), ncols):
            row.cells[j].text = ""
    doc.add_paragraph()


def _markdown_line_to_docx(doc: Document, line: str) -> None:
    """
    Append one line of markdown to the document as native Word content.
    Handles # ## ### headings, - / * bullet lists, **bold** / *italic* inline, and plain paragraphs.
    (Markdown tables are handled separately by block processing.)
    """
    stripped = line.strip()
    if not stripped:
        doc.add_paragraph()
        return
    if stripped.startswith("### "):
        doc.add_heading(stripped[4:].strip(), level=3)
    elif stripped.startswith("## "):
        doc.add_heading(stripped[3:].strip(), level=2)
    elif stripped.startswith("# "):
        doc.add_heading(stripped[2:].strip(), level=1)
    elif stripped.startswith("- ") or (stripped.startswith("* ") and not stripped.startswith("**")):
        bullet_text = stripped[2:].strip()
        _add_paragraph_with_inline_format(doc, bullet_text, style="List Bullet")
    else:
        _add_paragraph_with_inline_format(doc, stripped)


def _report_blocks(report_text: str) -> list[tuple[str, str | list[str]]]:
    """
    Split report text into blocks: ("paragraph", line) or ("table", list_of_lines).
    Consecutive markdown table rows are grouped into one table block.
    """
    lines = report_text.strip().split("\n")
    blocks = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if _is_markdown_table_row(line.strip()):
            table_lines = []
            while i < len(lines) and _is_markdown_table_row(lines[i].strip()):
                table_lines.append(lines[i])
                i += 1
            blocks.append(("table", table_lines))
            continue
        blocks.append(("paragraph", line))
        i += 1
    return blocks


def write_report_docx(report_text: str, output_path: Path | None = None) -> Path:
    """
    Write report_text to a Word document. Converts markdown (headings, bold, lists)
    to native Word formatting. If output_path is None, uses
    reports/mbta_commute_report_{timestamp}.docx under module dir.
    Returns the path where the file was saved.
    """
    if output_path is None:
        timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
        reports_dir = _app_dir / "reports"
        reports_dir.mkdir(parents=True, exist_ok=True)
        output_path = reports_dir / f"mbta_commute_report_{timestamp}.docx"
    doc = Document()
    # Title as heading and timestamp
    doc.add_heading("MBTA Red Line – Morning Commute Overview", level=0)
    doc.add_paragraph(datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC"))
    doc.add_paragraph()
    # Convert report into blocks (paragraphs vs markdown tables), then to native Word
    blocks = _report_blocks(report_text)
    for kind, content in blocks:
        if kind == "table":
            _add_markdown_table_to_docx(doc, content)
        else:
            _markdown_line_to_docx(doc, content)
    doc.save(str(output_path))
    return output_path


# 6. Main flow #################################################################


def main() -> None:
    """Fetch API data, build dataframes, send to Ollama Cloud, write .docx report."""
    if not MBTA_API_KEY:
        raise ValueError("MBTA_API_KEY is not set. Add it to .env.")
    if not OLLAMA_API_KEY:
        raise ValueError("OLLAMA_API_KEY is not set. Add it to .env.")

    # 1. Fetch all API data
    alerts_resp = fetch_alerts()
    predictions_alewife_resp = fetch_predictions_at_stop(STOP_ALEWIFE)
    predictions_central_sq_resp = fetch_predictions_at_stop(STOP_CENTRAL_SQ)
    vehicles_resp = fetch_vehicles()

    # Check for API errors
    for name, resp in [
        ("alerts", alerts_resp),
        ("predictions Alewife", predictions_alewife_resp),
        ("predictions Central Square", predictions_central_sq_resp),
        ("vehicles", vehicles_resp),
    ]:
        if resp.get("error"):
            raise ValueError(f"MBTA API error ({name}): {resp.get('message', 'Unknown')}")

    # 2. Build unfiltered dataframes
    df_alerts = build_alerts_df(alerts_resp)
    df_predictions_alewife = build_predictions_df(predictions_alewife_resp)
    df_predictions_central_sq = build_predictions_df(predictions_central_sq_resp)
    df_vehicles = build_vehicles_df(vehicles_resp)

    # 3. Format for Ollama
    formatted_data = format_data_for_ollama(
        df_alerts,
        df_predictions_alewife,
        df_predictions_central_sq,
        df_vehicles,
    )
    user_content = USER_PROMPT.strip() + "\n\n---\nData:\n" + formatted_data

    # 4. Call Ollama Cloud
    report_text = query_ollama_cloud(user_content)

    # 5. Write response to .docx
    out_path = write_report_docx(report_text)
    print(f"Report saved to: {out_path}")


if __name__ == "__main__":
    main()
